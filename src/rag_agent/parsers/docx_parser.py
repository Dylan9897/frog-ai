"""
Word 文档解析器
将 .docx 文件转换为图片
"""
from pathlib import Path
from typing import Dict, Any, Optional, List
from PIL import Image, ImageDraw, ImageFont

from .base_parser import BaseParser

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxParser(BaseParser):
    """Word 文档解析器"""
    
    def __init__(self, pages_dir: str, dpi: int = 150):
        """
        初始化 Word 解析器
        参数:
          - pages_dir: 页面图片存储目录
          - dpi: 图片分辨率（默认 150）
        """
        self.pages_dir = Path(pages_dir)
        self.pages_dir.mkdir(parents=True, exist_ok=True)
        self.dpi = dpi
        # 页面尺寸（A4 比例）
        self.page_width = int(8.27 * dpi)
        self.page_height = int(11.69 * dpi)
        self.margin = int(0.5 * dpi)
        self.line_height = int(0.2 * dpi)
        self.font_size = int(0.12 * dpi)
    
    def _get_font(self, size: Optional[int] = None):
        """获取字体"""
        try:
            if Path.exists(Path("C:/Windows/Fonts/msyh.ttc")):
                return ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", size or self.font_size)
            elif Path.exists(Path("C:/Windows/Fonts/simsun.ttc")):
                return ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", size or self.font_size)
            else:
                return ImageFont.load_default()
        except:
            return ImageFont.load_default()
    
    def _wrap_text(self, text: str, font, max_width: int) -> List[str]:
        """文本换行"""
        lines = []
        words = text.split('\n')
        for word_line in words:
            if not word_line:
                lines.append('')
                continue
            current_line = ''
            for char in word_line:
                test_line = current_line + char
                bbox = font.getbbox(test_line)
                if bbox[2] - bbox[0] <= max_width:
                    current_line = test_line
                else:
                    if current_line:
                        lines.append(current_line)
                    current_line = char
            if current_line:
                lines.append(current_line)
        return lines
    
    def _render_text_to_image(self, text: str) -> Image.Image:
        """将文本渲染为图片"""
        img = Image.new('RGB', (self.page_width, self.page_height), color='white')
        draw = ImageDraw.Draw(img)
        font = self._get_font()
        
        text_area_width = self.page_width - 2 * self.margin
        text_area_height = self.page_height - 2 * self.margin
        lines = self._wrap_text(text, font, text_area_width)
        
        y = self.margin
        for line in lines:
            if y + self.line_height > self.page_height - self.margin:
                break
            draw.text((self.margin, y), line, fill='black', font=font)
            y += self.line_height
        
        return img
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析 Word 文档"""
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "python-docx 未安装",
                "content": "",
                "pages": [],
                "metadata": {}
            }
        
        try:
            doc = DocxDocument(file_path)
            parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)
            content = "\n".join(parts)
            
            return {
                "success": True,
                "content": content,
                "pages": [{"page_number": 1, "text": content}],
                "metadata": {},
                "total_pages": 1
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "content": "",
                "pages": [],
                "metadata": {}
            }
    
    def convert_to_images(
        self, 
        file_path: str, 
        document_id: str,
        format: str = 'png'
    ) -> List[Dict[str, Any]]:
        """将 Word 文档转换为图片"""
        if not DOCX_AVAILABLE:
            print("[DocxParser] python-docx 未安装，无法转换 Word 文档")
            return []
        
        try:
            # 读取 Word 内容
            doc = DocxDocument(file_path)
            parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)
            content = "\n".join(parts)
            
            # 创建文档专用目录
            doc_pages_dir = self.pages_dir / document_id
            doc_pages_dir.mkdir(parents=True, exist_ok=True)
            
            pages_info = []
            
            # 计算分页
            font = self._get_font()
            text_area_width = self.page_width - 2 * self.margin
            lines = self._wrap_text(content, font, text_area_width)
            lines_per_page = (self.page_height - 2 * self.margin) // self.line_height
            total_pages = max(1, (len(lines) + lines_per_page - 1) // lines_per_page)
            
            # 分页渲染
            for page_num in range(total_pages):
                start_line = page_num * lines_per_page
                end_line = min(start_line + lines_per_page, len(lines))
                page_text = '\n'.join(lines[start_line:end_line])
                
                img = self._render_text_to_image(page_text)
                
                image_filename = f"page_{page_num + 1:04d}.{format}"
                image_path = doc_pages_dir / image_filename
                img.save(str(image_path), format=format.upper())
                
                pages_info.append({
                    "page_number": page_num + 1,
                    "image_path": str(image_path),
                    "width": img.width,
                    "height": img.height,
                    "format": format
                })
            
            return pages_info
        except Exception as e:
            print(f"[DocxParser] 转换图片失败: {e}")
            import traceback
            traceback.print_exc()
            return []
    
    def get_page_content(
        self, 
        file_path: str, 
        page_number: int,
        format: str = 'image/png'
    ) -> Optional[bytes]:
        """获取指定页面内容"""
        if not DOCX_AVAILABLE:
            return None
        
        try:
            doc = DocxDocument(file_path)
            parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)
            content = "\n".join(parts)
            
            if format.startswith('image/'):
                img = self._render_text_to_image(content)
                from io import BytesIO
                img_bytes = BytesIO()
                img.save(img_bytes, format=format.split('/')[1].upper())
                return img_bytes.getvalue()
            else:
                return content.encode('utf-8')
        except Exception as e:
            print(f"[DocxParser] 获取页面内容失败: {e}")
            return None
    
    def get_total_pages(self, file_path: str) -> int:
        """获取文档总页数"""
        if not DOCX_AVAILABLE:
            return 0
        
        try:
            doc = DocxDocument(file_path)
            parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)
            content = "\n".join(parts)
            
            font = self._get_font()
            text_area_width = self.page_width - 2 * self.margin
            lines = self._wrap_text(content, font, text_area_width)
            lines_per_page = (self.page_height - 2 * self.margin) // self.line_height
            return max(1, (len(lines) + lines_per_page - 1) // lines_per_page)
        except Exception as e:
            print(f"[DocxParser] 获取总页数失败: {e}")
            return 1


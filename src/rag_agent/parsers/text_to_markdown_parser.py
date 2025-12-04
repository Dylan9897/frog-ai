"""
文本转 Markdown 解析器
支持 doc/docx/txt/json/excel/md 等格式转换为 Markdown
"""
from pathlib import Path
from typing import Dict, Any, Optional, List
import json
import markdown
import re

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    import subprocess
    import os
    LIBREOFFICE_AVAILABLE = True
except ImportError:
    LIBREOFFICE_AVAILABLE = False


class TextToMarkdownParser:
    """文本转 Markdown 解析器"""
    
    @staticmethod
    def doc_to_docx(doc_path: str, output_dir: Path) -> Optional[str]:
        """将 .doc 文件转换为 .docx 文件"""
        if not LIBREOFFICE_AVAILABLE:
            return None
        
        try:
            # 查找 LibreOffice
            soffice_paths = [
                "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
                "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
                "soffice",
            ]
            
            soffice = None
            for path in soffice_paths:
                if path == "soffice" or os.path.exists(path):
                    try:
                        result = subprocess.run(
                            [path, "--version"],
                            capture_output=True,
                            timeout=5
                        )
                        if result.returncode == 0:
                            soffice = path
                            break
                    except:
                        continue
            
            if not soffice:
                print("[TextToMarkdownParser] 未找到 LibreOffice，无法转换 .doc 文件")
                return None
            
            # 转换为 docx
            output_dir.mkdir(parents=True, exist_ok=True)
            docx_path = output_dir / (Path(doc_path).stem + ".docx")
            
            cmd = [
                soffice,
                "--headless",
                "--convert-to", "docx",
                "--outdir", str(output_dir),
                str(Path(doc_path).absolute())
            ]
            
            result = subprocess.run(cmd, capture_output=True, timeout=60)
            
            if result.returncode == 0 and docx_path.exists():
                return str(docx_path)
            return None
        except Exception as e:
            print(f"[TextToMarkdownParser] DOC 转 DOCX 失败: {e}")
            return None
    
    @staticmethod
    def docx_to_markdown(file_path: str) -> str:
        """将 DOCX 文件转换为 Markdown"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安装")
        
        doc = DocxDocument(file_path)
        markdown_parts = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                markdown_parts.append("")
                continue
            
            # 检查段落样式
            style_name = para.style.name.lower() if para.style else ""
            
            # 标题处理
            if "heading" in style_name or "title" in style_name:
                level = 1
                if "heading 1" in style_name or "title" in style_name:
                    level = 1
                elif "heading 2" in style_name:
                    level = 2
                elif "heading 3" in style_name:
                    level = 3
                else:
                    level = 1
                markdown_parts.append(f"{'#' * level} {text}")
            else:
                # 处理加粗、斜体等格式
                formatted_text = ""
                for run in para.runs:
                    run_text = run.text
                    if run.bold:
                        run_text = f"**{run_text}**"
                    if run.italic:
                        run_text = f"*{run_text}*"
                    formatted_text += run_text
                
                if formatted_text:
                    markdown_parts.append(formatted_text)
                else:
                    markdown_parts.append(text)
        
        # 处理表格
        for table in doc.tables:
            markdown_parts.append("")
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
                markdown_parts.append("| " + " | ".join(cells) + " |")
                if i == 0:  # 表头分隔符
                    markdown_parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
            markdown_parts.append("")
        
        return "\n".join(markdown_parts)
    
    @staticmethod
    def txt_to_markdown(file_path: str) -> str:
        """将 TXT 文件转换为 Markdown"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # 简单的文本转 Markdown（保持原样，可以添加一些格式化）
        # 检测是否是代码（简单判断）
        if any(keyword in content.lower() for keyword in ['def ', 'function', 'import ', 'class ', '<?php', '<html']):
            return f"```\n{content}\n```"
        
        return content
    
    @staticmethod
    def json_to_markdown(file_path: str) -> str:
        """将 JSON 文件转换为 Markdown"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 格式化为 Markdown 代码块
        json_str = json.dumps(data, ensure_ascii=False, indent=2)
        return f"```json\n{json_str}\n```"
    
    @staticmethod
    def excel_to_markdown(file_path: str) -> str:
        """将 Excel 文件转换为 Markdown"""
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl 未安装")
        
        wb = load_workbook(file_path, read_only=True, data_only=True)
        markdown_parts = []
        
        for sheet_name in wb.sheetnames:
            markdown_parts.append(f"## 工作表: {sheet_name}")
            markdown_parts.append("")
            
            sheet = wb[sheet_name]
            rows_data = []
            
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                if any(cell.strip() for cell in row_data):  # 跳过空行
                    rows_data.append(row_data)
            
            if rows_data:
                # 创建表格
                for i, row in enumerate(rows_data):
                    cells = [cell.replace("|", "\\|") for cell in row]
                    markdown_parts.append("| " + " | ".join(cells) + " |")
                    if i == 0:  # 表头分隔符
                        markdown_parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
                markdown_parts.append("")
        
        wb.close()
        return "\n".join(markdown_parts)
    
    @staticmethod
    def md_to_markdown(file_path: str) -> str:
        """Markdown 文件直接读取"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @staticmethod
    def parse_to_markdown(file_path: str, file_ext: str, temp_dir: Optional[Path] = None) -> str:
        """根据文件类型解析为 Markdown"""
        file_ext = file_ext.lower().lstrip('.')
        
        # 处理 .doc 文件（需要先转换为 .docx）
        if file_ext == 'doc':
            if temp_dir is None:
                temp_dir = Path(file_path).parent / "temp"
            temp_dir = Path(temp_dir)
            temp_dir.mkdir(parents=True, exist_ok=True)
            
            docx_path = TextToMarkdownParser.doc_to_docx(file_path, temp_dir)
            if docx_path:
                return TextToMarkdownParser.docx_to_markdown(docx_path)
            else:
                raise Exception("无法将 .doc 文件转换为 .docx")
        
        # 处理其他格式
        if file_ext == 'docx':
            return TextToMarkdownParser.docx_to_markdown(file_path)
        elif file_ext == 'txt':
            return TextToMarkdownParser.txt_to_markdown(file_path)
        elif file_ext == 'json':
            return TextToMarkdownParser.json_to_markdown(file_path)
        elif file_ext in ['xlsx', 'xls']:
            return TextToMarkdownParser.excel_to_markdown(file_path)
        elif file_ext in ['md', 'markdown']:
            return TextToMarkdownParser.md_to_markdown(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")

